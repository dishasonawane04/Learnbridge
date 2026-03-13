import os
from langchain_core.documents import Document

def load_document(file_path):
    """
    Load generic document based on file extension.
    Supports PDF, DOCX, TXT, and Images.
    Returns a list of langchain Document objects.
    """
    ext = os.path.splitext(file_path)[1].lower()
    pages = []
    
    if ext == '.pdf':
        import fitz
        with fitz.open(file_path) as doc:
            for i, page in enumerate(doc):
                text = page.get_text()
                if text.strip():
                    pages.append(Document(page_content=text, metadata={"source": file_path, "page": i+1}))
    elif ext in ['.docx', '.doc']:
        # Fallback to simple docx parsing
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            text = '\n'.join(full_text)
            if text.strip():
                pages.append(Document(page_content=text, metadata={"source": file_path, "page": 1}))
        except ImportError:
            pass
    elif ext == '.txt':
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
            if text.strip():
                pages.append(Document(page_content=text, metadata={"source": file_path, "page": 1}))
    elif ext in ['.png', '.jpg', '.jpeg']:
        from course.utils.extraction import extract_text_from_image
        text = extract_text_from_image(file_path)
        if text.strip() and not text.startswith("Image OCR Failed"):
            pages.append(Document(page_content=text, metadata={"source": file_path, "page": 1}))
            
    return pages
