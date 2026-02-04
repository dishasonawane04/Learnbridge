from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse, HttpResponseRedirect
from django.urls import reverse
from django.contrib import messages
from .models import LetterRequest
from .forms import LetterRequestForm
from .ai_logic import generate_recommendation_letter
from accounts.decorators import teacher_required
from analytics.models import ActivityLog
import io
from django.db.models import Avg
from course.models import Course

# Document generation imports
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Pt
import datetime

@teacher_required
def letter_dashboard(request):
    """
    Displays the dashboard for the LOR Generator, including history and a 'Create New' button.
    """
    letters = LetterRequest.objects.all().order_by('-created_at')
    return render(request, 'letter_of_recommendation_generator/dashboard.html', {'letters': letters})

@teacher_required
def create_letter(request):
    """
    Handles the multi-step form submission. (Currently simplified to one page for MVP, can be split with JS).
    """
    if request.method == 'POST':
        form = LetterRequestForm(request.POST)
        if form.is_valid():
            letter_request = form.save(commit=False)
            
            # Generate the letter using AI
            generated_content = generate_recommendation_letter(form.cleaned_data)
            letter_request.generated_letter = generated_content
            
            letter_request.save()

            # --- ANALYTICS LOGGING ---
            if request.user.is_authenticated:
                ActivityLog.objects.create(
                    user=request.user,
                    app_name='lor',
                    activity_type='letter_generated',
                    topic=f"LOR for {letter_request.student_name}",
                    metadata={'purpose': letter_request.purpose}
                )
            # -------------------------

            return redirect('letter_of_recommendation_generator:preview_letter', pk=letter_request.pk)
    else:
        initial_data = {}
        course_id = request.GET.get('course_id')
        student_id = request.GET.get('student_id')
        
        if course_id and student_id:
            # Auto-fill Logic
            # 1. Get Course & Student info (mock or real if User model had name)
            # student = User.objects.get(id=student_id) ...
            
            # 2. Analyze Performance
            # Calculate avg quiz score, decks created, etc.
            logs = ActivityLog.objects.filter(user_id=student_id)
            
            quiz_score_avg = logs.filter(activity_type='quiz_completed').aggregate(Avg('score'))['score__avg']
            decks_count = logs.filter(activity_type='deck_generated').count()
            
            perf_text = f"Analyzed Data:\n- Flashcard Decks Created: {decks_count}\n"
            if quiz_score_avg:
                perf_text += f"- Average Quiz Score: {quiz_score_avg:.1f}%\n"
                
            initial_data['academic_performance'] = perf_text
            initial_data['purpose'] = 'higher_studies' # Default
            
        form = LetterRequestForm(initial=initial_data)
    
    return render(request, 'letter_of_recommendation_generator/form.html', {'form': form})

@teacher_required
def preview_letter(request, pk):
    """
    Allows the user to view and edit the generated letter.
    """
    letter = get_object_or_404(LetterRequest, pk=pk)
    
    if request.method == 'POST':
        # Save changes
        updated_content = request.POST.get('letter_content')
        if updated_content:
            letter.generated_letter = updated_content
            letter.save()
            messages.success(request, "Letter updated successfully.")
            return redirect('letter_of_recommendation_generator:preview_letter', pk=pk)
            
    return render(request, 'letter_of_recommendation_generator/preview.html', {'letter': letter})

@teacher_required
def delete_letter(request, pk):
    letter = get_object_or_404(LetterRequest, pk=pk)
    letter.delete()
    messages.success(request, "Letter deleted.")
    return redirect('letter_of_recommendation_generator:letter_dashboard')

@teacher_required
def download_letter_pdf(request, pk):
    """
    Downloads the letter as a PDF using xhtml2pdf.
    """
    letter = get_object_or_404(LetterRequest, pk=pk)
    
    # Simple HTML template for the PDF
    html_content = f"""
    <html>
    <head>
        <style>
            @page {{
                size: letter;
                margin: 2cm;
            }}
            body {{
                font-family: "Times New Roman", Times, serif;
                font-size: 12pt;
                line-height: 1.5;
            }}
            .header {{
                text-align: center;
                margin-bottom: 2cm;
                font-weight: bold;
                font-size: 16pt;
            }}
            .date {{
                text-align: right;
                margin-bottom: 1cm;
            }}
            .content {{
                text-align: justify;
                white-space: pre-line;
            }}
            .footer {{
                margin-top: 2cm;
                font-style: italic;
            }}
        </style>
    </head>
    <body>
        <div class="header">LETTER OF RECOMMENDATION</div>
        <div class="date">{datetime.date.today().strftime("%B %d, %Y")}</div>
        <div class="content">{letter.generated_letter}</div>
        <div class="footer">
            <p>Generated via LearnBridge Academic Tools</p>
        </div>
    </body>
    </html>
    """
    
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{letter.student_name.replace(" ", "_")}_LOR.pdf"'
    
    pisa_status = pisa.CreatePDF(html_content, dest=response)
    
    if pisa_status.err:
        return HttpResponse('We had some errors <pre>' + html_content + '</pre>')
    return response

@teacher_required
def download_letter_docx(request, pk):
    """
    Downloads the letter as a Word Document using python-docx.
    """
    letter = get_object_or_404(LetterRequest, pk=pk)
    
    document = Document()
    
    # Add Title
    document.add_heading('Letter of Recommendation', 0)
    
    # Add Date
    p = document.add_paragraph(datetime.date.today().strftime("%B %d, %Y"))
    p.alignment = 2  # Right alignment
    
    # Add Content
    # Split by newlines to create proper paragraphs
    paragraphs = letter.generated_letter.split('\n')
    for para in paragraphs:
        if para.strip():
            p = document.add_paragraph(para.strip())
            p.style.font.name = 'Times New Roman'
            p.style.font.size = Pt(12)
            
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{letter.student_name.replace(" ", "_")}_LOR.docx"'
    
    document.save(response)
    return response
